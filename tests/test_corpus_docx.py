# Corpus tests: DOCX structures.
#   - 3-level nested tables             - gridSpan + vMerge merged cells
#   - cell containing hyperlink+image   - multiple hyperlinks in a paragraph
#   - bold-formatted hyperlink          - multi-level numbered lists
#   - w:tab / w:br round-trip           - >256-token paragraph (chunk re-join)
#   - CJK+Latin mixed text              - textbox via raw XML
#
# Run from the repo root:
#   python tests/test_corpus_docx.py
import json
import os
import re
import sys
import zipfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from tests.corpus_common import T, check, fake_translate, run, work_dirs

WORK_DIR, TEMP_DIR, RESULT_DIR = work_dirs("docx")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def add_hyperlink(paragraph, url, text, bold=False):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run_el = OxmlElement("w:r")
    if bold:
        rpr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rpr.append(b)
        run_el.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run_el.append(t)
    hl.append(run_el)
    paragraph._p.append(hl)


def make_png(path):
    from PIL import Image
    Image.new("RGB", (50, 30), (40, 160, 90)).save(path)
    return path


def doc_xml(path):
    with zipfile.ZipFile(path) as z:
        return z.read("word/document.xml").decode("utf-8")


def paragraph_texts(path):
    """Full text of each w:p in document order (hyperlink runs included,
    tabs as \\t)."""
    from lxml import etree
    with zipfile.ZipFile(path) as z:
        tree = etree.fromstring(z.read("word/document.xml"))
    texts = []
    for p in tree.iter(f"{{{W_NS}}}p"):
        parts = []
        for node in p.iter():
            if node.tag == f"{{{W_NS}}}t":
                parts.append(node.text or "")
            elif node.tag == f"{{{W_NS}}}tab":
                parts.append("\t")
        texts.append("".join(parts))
    return texts


def translate_docx(src):
    from pipeline.word_translation_pipeline import (
        extract_word_content_to_json, write_translated_content_to_word)
    src_json = extract_word_content_to_json(src, TEMP_DIR)
    dst_json = fake_translate(src_json)
    return write_translated_content_to_word(src, src_json, dst_json, TEMP_DIR, RESULT_DIR,
                                            bilingual_mode=False, src_lang="en", dst_lang="ja")


# --------------------------------------------------------------- tests ----
def test_nested_tables_and_merges():
    print("DOCX: 3-level nested tables + gridSpan/vMerge merged cells")
    from docx import Document

    src = os.path.join(WORK_DIR, "tables.docx")
    doc = Document()

    # 3-level nesting: outer > middle (in outer cell) > inner (in middle cell)
    outer = doc.add_table(rows=2, cols=2, style="Table Grid")
    outer.cell(0, 0).text = "Level one outer cell"
    outer.cell(1, 1).text = "Level one bottom cell"
    middle = outer.cell(0, 1).add_table(rows=1, cols=2)
    middle.cell(0, 0).text = "Level two middle cell"
    inner = middle.cell(0, 1).add_table(rows=1, cols=1)
    inner.cell(0, 0).text = "Level three innermost cell"

    # merged cells: gridSpan (horizontal) + vMerge (vertical)
    merged = doc.add_table(rows=3, cols=3, style="Table Grid")
    merged.cell(0, 0).merge(merged.cell(0, 2))   # gridSpan across the top row
    merged.cell(1, 0).merge(merged.cell(2, 0))   # vMerge down the left column
    merged.cell(0, 0).text = "Horizontally merged header"
    merged.cell(1, 0).text = "Vertically merged label"
    merged.cell(1, 1).text = "Plain middle cell"
    merged.cell(2, 2).text = "Plain corner cell"
    doc.save(src)

    out = translate_docx(src)
    xml = doc_xml(out)
    joined = "\n".join(paragraph_texts(out))

    check("all three nesting levels translated",
          all(T + s in joined for s in ("Level one outer cell", "Level two middle cell",
                                        "Level three innermost cell", "Level one bottom cell")),
          joined)
    check("all 4 w:tbl elements survive (3-level nest + merged table)",
          xml.count("<w:tbl>") == 4, f"found {xml.count('<w:tbl>')}")
    # the innermost text must sit under three stacked w:tbl ancestors
    from lxml import etree
    tree = etree.fromstring(xml.encode("utf-8"))
    innermost = next(t for t in tree.iter(f"{{{W_NS}}}t")
                     if "Level three innermost cell" in (t.text or ""))
    depth = sum(1 for a in innermost.iterancestors() if a.tag == f"{{{W_NS}}}tbl")
    check("innermost cell still nested 3 tables deep", depth == 3, f"depth={depth}")
    check("gridSpan survives", re.search(r'<w:gridSpan w:val="3"\s*/>', xml) is not None, xml[:800])
    check("vMerge survives (restart + continue)",
          'w:vMerge w:val="restart"' in xml and "<w:vMerge/>" in xml
          or xml.count("vMerge") >= 2, xml[:800])
    check("merged cell texts translated",
          T + "Horizontally merged header" in joined
          and T + "Vertically merged label" in joined, joined)
    check("plain cells translated",
          T + "Plain middle cell" in joined and T + "Plain corner cell" in joined, joined)


def test_hyperlinks_and_cell_with_image():
    print("DOCX: multiple/bold hyperlinks + table cell with hyperlink and image")
    from docx import Document
    from docx.shared import Inches

    src = os.path.join(WORK_DIR, "links.docx")
    doc = Document()

    # multiple hyperlinks in ONE paragraph
    p1 = doc.add_paragraph("Leading text before ")
    add_hyperlink(p1, "https://example.com/first", "first linked words")
    p1.add_run(" middle joining text ")
    add_hyperlink(p1, "https://example.com/second", "second linked words")
    p1.add_run(" trailing closing text.")

    # bold hyperlink
    p2 = doc.add_paragraph("A sentence holding ")
    add_hyperlink(p2, "https://example.com/bold", "a bold emphasized link", bold=True)
    p2.add_run(" in the middle of it.")

    # table cell containing hyperlink + image
    table = doc.add_table(rows=1, cols=1, style="Table Grid")
    cell_p = table.cell(0, 0).paragraphs[0]
    cell_p.add_run("Cell text before link ")
    add_hyperlink(cell_p, "https://example.com/cell", "cell anchor words")
    cell_p.add_run(" then an image ")
    cell_p.add_run().add_picture(make_png(os.path.join(WORK_DIR, "pic.png")),
                                 width=Inches(0.3))
    cell_p.add_run(" and tail words.")
    doc.save(src)

    out = translate_docx(src)
    xml = doc_xml(out)
    paras = paragraph_texts(out)
    with zipfile.ZipFile(out) as z:
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")

    from lxml import etree
    tree = etree.fromstring(xml.encode("utf-8"))
    link_texts = ["".join(t.text or "" for t in hl.iter(f"{{{W_NS}}}t"))
                  for hl in tree.iter(f"{{{W_NS}}}hyperlink")]

    check("all four hyperlink elements survive", len(link_texts) == 4, str(link_texts))
    check("all hyperlink targets survive in rels",
          all(u in rels for u in ("example.com/first", "example.com/second",
                                  "example.com/bold", "example.com/cell")), rels)
    check("link texts stayed inside their hyperlink elements",
          any("first linked words" in lt for lt in link_texts)
          and any("second linked words" in lt for lt in link_texts)
          and any("cell anchor words" in lt for lt in link_texts), str(link_texts))

    two_link_para = next((p for p in paras if "middle joining text" in p), "")
    order = [two_link_para.find("Leading text"), two_link_para.find("first linked words"),
             two_link_para.find("middle joining text"),
             two_link_para.find("second linked words"), two_link_para.find("trailing closing")]
    check("two-link paragraph order preserved and translated",
          T in two_link_para and all(i >= 0 for i in order) and order == sorted(order),
          repr(two_link_para))

    bold_hl = next((hl for hl in tree.iter(f"{{{W_NS}}}hyperlink")
                    if "bold emphasized link" in
                    "".join(t.text or "" for t in hl.iter(f"{{{W_NS}}}t"))), None)
    check("bold formatting survives on the hyperlink run",
          bold_hl is not None and bold_hl.find(f".//{{{W_NS}}}b") is not None,
          etree.tostring(bold_hl, encoding="unicode") if bold_hl is not None else "link lost")

    cell_para = next((p for p in paras if "cell anchor words" in p), "")
    check("cell paragraph translated with link+image and text order kept",
          T in cell_para
          and cell_para.find("Cell text before") < cell_para.find("cell anchor words")
          < cell_para.find("then an image") < cell_para.find("and tail words"),
          repr(cell_para))
    check("image drawing survives inside the table cell",
          "<w:drawing>" in xml[xml.find("<w:tbl>"):], xml[xml.find("<w:tbl>"):][:600])


def test_multilevel_numbered_lists():
    print("DOCX: multi-level numbered lists keep numbering and level")
    from docx import Document

    src = os.path.join(WORK_DIR, "lists.docx")
    doc = Document()
    doc.add_paragraph("Top level first entry", style="List Number")
    doc.add_paragraph("Second level child entry", style="List Number 2")
    doc.add_paragraph("Third level grandchild entry", style="List Number 3")
    doc.add_paragraph("Top level second entry", style="List Number")
    doc.save(src)

    out = translate_docx(src)
    xml = doc_xml(out)
    joined = "\n".join(paragraph_texts(out))

    check("all list entries translated",
          all(T + s in joined for s in ("Top level first entry", "Second level child entry",
                                        "Third level grandchild entry",
                                        "Top level second entry")), joined)
    check("list styles survive on the paragraphs",
          all(f'w:val="{s}"' in xml for s in ("ListNumber", "ListNumber2", "ListNumber3")),
          xml[:1200])
    order = [joined.find(s) for s in ("Top level first entry", "Second level child entry",
                                      "Third level grandchild entry", "Top level second entry")]
    check("list order preserved", all(i >= 0 for i in order) and order == sorted(order),
          str(order))


def test_tab_br_roundtrip():
    print("DOCX: w:tab and w:br round-trip as real elements")
    from docx import Document
    from docx.oxml import OxmlElement

    src = os.path.join(WORK_DIR, "tabbr.docx")
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Column one header")
    r1._r.append(OxmlElement("w:tab"))
    r2 = p.add_run("Column two header")
    r2._r.append(OxmlElement("w:br"))
    p.add_run("Second visual line text")
    doc.save(src)

    out = translate_docx(src)
    xml = doc_xml(out)
    paras = paragraph_texts(out)
    para = next((p for p in paras if "Column one header" in p), "")

    check("paragraph translated", T in para, repr(para))
    check("tab restored as a real w:tab element (not a literal character)",
          "<w:tab/>" in xml and not re.search(r"<w:t[^>]*>[^<]*\t", xml), xml[:1000])
    check("break restored as a real w:br element", "<w:br/>" in xml, xml[:1000])
    check("no newline markers leaked", "␊" not in xml and "␍" not in xml, xml[:1000])
    check("text segments in order around tab",
          para.find("Column one header") < para.find("\t") < para.find("Column two header"),
          repr(para))
    check("post-break text present", "Second visual line text" in "".join(paras),
          str(paras))


def test_long_paragraph_chunk_rejoin():
    print("DOCX: >256-token paragraphs split into chunks and re-joined without glue")
    from docx import Document
    from pipeline.word_translation_pipeline import (
        extract_word_content_to_json, write_translated_content_to_word)
    from textProcessing.text_separator import (
        deduplicate_translation_content, create_deduped_json_for_translation,
        split_text_by_token_limit, restore_translations_from_deduped)

    latin = " ".join(f"Sentence number {i} describes the architecture of the translation "
                     "system in detail." for i in range(1, 61))
    cjk = "".join(f"第{i}句は翻訳システムの構造を詳細に説明する長い文章です。" for i in range(1, 41))

    src = os.path.join(WORK_DIR, "long.docx")
    doc = Document()
    doc.add_paragraph(latin)
    doc.add_paragraph(cjk)
    doc.add_paragraph("Short trailing paragraph.")
    doc.save(src)

    src_json = extract_word_content_to_json(src, TEMP_DIR)

    # Full dedupe -> split -> fake-translate-each-chunk -> restore chain,
    # exactly as base_translator drives it (max 256 tokens per chunk)
    folder = os.path.dirname(src_json)
    deduped_path = os.path.join(folder, "src_deduped.json")
    deduped_data, src_to_deduped = deduplicate_translation_content(src_json)
    create_deduped_json_for_translation(deduped_data, deduped_path)
    split_path = split_text_by_token_limit(deduped_path, max_tokens=256)

    with open(split_path, encoding="utf-8") as f:
        split_items = json.load(f)
    latin_chunks = [i for i in split_items if "Sentence number" in i["value"]]
    cjk_chunks = [i for i in split_items if "翻訳システム" in i["value"]]
    check("latin paragraph split into multiple chunks", len(latin_chunks) >= 2,
          f"{len(latin_chunks)} chunks")
    check("cjk paragraph split into multiple chunks", len(cjk_chunks) >= 2,
          f"{len(cjk_chunks)} chunks")

    dst_split_path = os.path.join(folder, "dst_translated_split.json")
    with open(dst_split_path, "w", encoding="utf-8") as f:
        json.dump([{"count_split": i["count_split"], "original": i["value"],
                    "translated": T + i["value"]} for i in split_items],
                  f, ensure_ascii=False, indent=2)

    dst_json = restore_translations_from_deduped(dst_split_path, src_to_deduped, src_json)
    out = write_translated_content_to_word(src, src_json, dst_json, TEMP_DIR, RESULT_DIR,
                                           bilingual_mode=False, src_lang="en", dst_lang="ja")

    paras = paragraph_texts(out)
    latin_para = next((p for p in paras if "Sentence number 1 " in p), "")
    cjk_para = next((p for p in paras if "第1句" in p), "")

    check("every latin sentence survived the chunk round-trip",
          all(f"Sentence number {i} " in latin_para
              or latin_para.endswith(f"Sentence number {i} describes the architecture "
                                     "of the translation system in detail.")
              for i in (1, 17, 33, 60)), latin_para[:300])
    check("each latin chunk carries its own [T] marker",
          latin_para.count(T) == len(latin_chunks),
          f"{latin_para.count(T)} markers vs {len(latin_chunks)} chunks")
    check("no word glue at latin chunk joins (space kept before each [T])",
          not re.search(r"[^\s]\[T\]", latin_para.replace("。[T]", "")), latin_para[:400])
    check("every cjk sentence survived", all(f"第{i}句" in cjk_para for i in (1, 20, 40)),
          cjk_para[:200])
    check("no space inserted at cjk chunk joins",
          "。[T]" in cjk_para and "。 [T]" not in cjk_para, cjk_para[:400])
    check("short paragraph still translated once",
          any(p == T + "Short trailing paragraph." for p in paras), str(paras[-3:]))


def test_cjk_latin_mixed():
    print("DOCX: CJK + Latin mixed paragraph")
    from docx import Document

    src = os.path.join(WORK_DIR, "mixed.docx")
    doc = Document()
    doc.add_paragraph("これは日本語 mixed with English words と中文字符 in one sentence.")
    doc.add_paragraph("Pure ASCII paragraph for contrast.")
    doc.save(src)

    out = translate_docx(src)
    joined = "\n".join(paragraph_texts(out))
    check("mixed CJK/Latin paragraph translated intact",
          T + "これは日本語 mixed with English words と中文字符 in one sentence." in joined,
          joined)
    check("ascii paragraph translated", T + "Pure ASCII paragraph for contrast." in joined,
          joined)


TEXTBOX_XML = """\
<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
     xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
     xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
     xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
     xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
  <w:r>
    <mc:AlternateContent>
      <mc:Choice Requires="wps">
        <w:drawing>
          <wp:inline distT="0" distB="0" distL="0" distR="0">
            <wp:extent cx="2286000" cy="571500"/>
            <wp:docPr id="7" name="TextBox 7"/>
            <a:graphic>
              <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                <wps:wsp>
                  <wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="2286000" cy="571500"/></a:xfrm>
                    <a:prstGeom prst="rect"><a:avLst/></a:prstGeom></wps:spPr>
                  <wps:txbx>
                    <w:txbxContent>
                      <w:p><w:r><w:t>Floating textbox message line</w:t></w:r></w:p>
                    </w:txbxContent>
                  </wps:txbx>
                  <wps:bodyPr rot="0" vert="horz"/>
                </wps:wsp>
              </a:graphicData>
            </a:graphic>
          </wp:inline>
        </w:drawing>
      </mc:Choice>
    </mc:AlternateContent>
  </w:r>
</w:p>
"""


def test_textbox_raw_xml():
    print("DOCX: wps textbox content (raw XML) translated in place")
    from docx import Document
    from lxml import etree

    src = os.path.join(WORK_DIR, "textbox.docx")
    doc = Document()
    doc.add_paragraph("Body paragraph beside the textbox.")
    doc._body._element.append(etree.fromstring(TEXTBOX_XML.encode("utf-8")))
    doc.save(src)

    from pipeline.word_translation_pipeline import extract_word_content_to_json
    src_json = extract_word_content_to_json(src, TEMP_DIR)
    with open(src_json, encoding="utf-8") as f:
        values = [i["value"] for i in json.load(f)]
    check("textbox text extracted for translation",
          any("Floating textbox message line" in v for v in values), str(values))

    out = translate_docx(src)
    xml = doc_xml(out)
    check("textbox structure survives", "<wps:txbx>" in xml and "<w:txbxContent>" in xml, xml)
    tb = xml[xml.find("<w:txbxContent>"):xml.find("</w:txbxContent>")]
    check("textbox text translated inside txbxContent",
          T + "Floating textbox message line" in tb, tb)
    check("body paragraph translated",
          T + "Body paragraph beside the textbox." in xml, xml[:600])


if __name__ == "__main__":
    run([test_nested_tables_and_merges, test_hyperlinks_and_cell_with_image,
         test_multilevel_numbered_lists, test_tab_br_roundtrip,
         test_long_paragraph_chunk_rejoin, test_cjk_latin_mixed, test_textbox_raw_xml])
