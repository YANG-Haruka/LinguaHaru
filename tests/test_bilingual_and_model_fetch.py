# Tests for bilingual output modes (SRT/VTT/TXT/MD + existing Excel/Word)
# and the online model list auto-fetch (OpenAI-compatible GET /models).
#
# Run from the repo root:
#   python tests/test_bilingual_and_model_fetch.py
import json
import os
import shutil
import sys
import zipfile

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, REPO_ROOT)
os.chdir(REPO_ROOT)

WORK_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_bilingual_work")
TEMP_DIR = os.path.join(WORK_DIR, "temp")
RESULT_DIR = os.path.join(WORK_DIR, "result")

T = "[T]"  # fake-translation marker

PASSED, FAILED = [], []


def check(name, cond, detail=""):
    (PASSED if cond else FAILED).append(name)
    print(f"  [{'PASS' if cond else 'FAIL'}] {name}" + (f" -> {detail}" if detail and not cond else ""))
    return bool(cond)


def fake_translate(src_json_path):
    """Produce dst_translated.json next to src.json: prefix every value with [T]."""
    with open(src_json_path, encoding="utf-8") as f:
        data = json.load(f)
    out = []
    for item in data:
        out.append({
            "count_src": item["count_src"],
            "type": item.get("type", "text"),
            "original": item["value"],
            "translated": T + item["value"],
        })
    dst_path = os.path.join(os.path.dirname(src_json_path), "dst_translated.json")
    with open(dst_path, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
    return dst_path


# --------------------------------------------------------------- subtitles --
def test_srt_bilingual():
    print("SRT bilingual: translation first, original below, timestamps intact")
    from pipeline.subtitle_translation_pipeline import (
        extract_srt_content_to_json, write_translated_content_to_srt)

    src = os.path.join(WORK_DIR, "bi_test.srt")
    with open(src, "w", encoding="utf-8") as f:
        f.write("1\n00:00:01,500 --> 00:00:03,000\nHello there my friend\n\n"
                "2\n00:00:04,000 --> 00:00:05,250\nSecond subtitle line\n\n")

    src_json = extract_srt_content_to_json(src, TEMP_DIR)
    dst_json = fake_translate(src_json)
    out = write_translated_content_to_srt(src, src_json, dst_json, RESULT_DIR,
                                          src_lang="en", dst_lang="ja", bilingual_mode=True)
    with open(out, encoding="utf-8") as f:
        content = f.read()

    check("cue 1 has translation then original",
          f"{T}Hello there my friend\nHello there my friend" in content, content)
    check("cue 2 has translation then original",
          f"{T}Second subtitle line\nSecond subtitle line" in content, content)
    check("still exactly two timestamp lines", content.count("-->") == 2, content)
    check("original appears exactly once per cue (no triple lines)",
          content.count("Hello there my friend") == 2, content)

    # Monolingual output unchanged by the new parameter (default False)
    out_mono = write_translated_content_to_srt(src, src_json, dst_json, RESULT_DIR,
                                               src_lang="en", dst_lang="fr")
    with open(out_mono, encoding="utf-8") as f:
        mono = f.read()
    check("default stays monolingual", mono.count("Hello there my friend") == 1, mono)


def test_vtt_bilingual():
    print("VTT bilingual: translation first, original below, header intact")
    from pipeline.subtitle_formats_pipeline import (
        extract_vtt_content_to_json, write_translated_content_to_vtt)

    src = os.path.join(WORK_DIR, "bi_test.vtt")
    with open(src, "w", encoding="utf-8") as f:
        f.write("WEBVTT\n\n"
                "00:00:01.000 --> 00:00:03.000\nHello my friend\n\n"
                "00:00:04.000 --> 00:00:06.000\nSecond line here\n")

    src_json = extract_vtt_content_to_json(src, TEMP_DIR)
    dst_json = fake_translate(src_json)
    out = write_translated_content_to_vtt(src, src_json, dst_json, TEMP_DIR, RESULT_DIR,
                                          src_lang="en", dst_lang="ja", bilingual_mode=True)
    with open(out, encoding="utf-8") as f:
        content = f.read()

    check("WEBVTT header preserved", content.startswith("WEBVTT"), content[:30])
    check("cue 1 bilingual", f"{T}Hello my friend\nHello my friend" in content, content)
    check("cue 2 bilingual", f"{T}Second line here\nSecond line here" in content, content)
    check("timestamps intact", content.count("-->") == 2, content)


# --------------------------------------------------------------------- txt --
def test_txt_bilingual():
    print("TXT bilingual: translated line followed by original, blanks preserved")
    from pipeline.txt_translation_pipeline import (
        extract_txt_content_to_json, write_translated_content_to_txt)

    src = os.path.join(WORK_DIR, "bi_test.txt")
    with open(src, "w", encoding="utf-8") as f:
        f.write("Document title here\n"
                "\n"
                "    Indented paragraph line\n"
                "\n"
                "12345\n"          # numbers only: not translated, stays single
                "Final closing line\n")

    src_json = extract_txt_content_to_json(src, TEMP_DIR)
    dst_json = fake_translate(src_json)
    out = write_translated_content_to_txt(src, src_json, dst_json, TEMP_DIR, RESULT_DIR,
                                          src_lang="en", dst_lang="ja", bilingual_mode=True)
    with open(out, encoding="utf-8") as f:
        result = f.read()

    expected = ("[T]Document title here\n"
                "Document title here\n"
                "\n"
                "    [T]Indented paragraph line\n"
                "    Indented paragraph line\n"
                "\n"
                "12345\n"
                "[T]Final closing line\n"
                "Final closing line\n")
    check("bilingual line structure exact", result == expected, repr(result))


# ---------------------------------------------------------------------- md --
def test_md_bilingual():
    print("MD bilingual: blockquote originals for text lines only, tables/code untouched")
    from pipeline.md_translation_pipeline import (
        extract_md_content_to_json, write_translated_content_to_md)

    src = os.path.join(WORK_DIR, "bi_test.md")
    with open(src, "w", encoding="utf-8") as f:
        f.write("# Heading line for test\n"
                "\n"
                "Some plain paragraph text\n"
                "\n"
                "| Alpha cell | Beta cell |\n"
                "| --- | --- |\n"
                "| Gamma data | Delta data |\n"
                "\n"
                "```\n"
                "code_line = 1\n"
                "```\n")

    src_json = extract_md_content_to_json(src, TEMP_DIR)
    dst_json = fake_translate(src_json)
    out = write_translated_content_to_md(src, src_json, dst_json, TEMP_DIR, RESULT_DIR,
                                         src_lang="en", dst_lang="ja", bilingual_mode=True)
    with open(out, encoding="utf-8") as f:
        content = f.read()

    check("paragraph gets blockquote original",
          f"{T}Some plain paragraph text\n> Some plain paragraph text" in content, content)
    check("heading gets blockquote original",
          f"{T}# Heading line for test\n> # Heading line for test" in content, content)
    check("no blockquote injected inside pipe table", "> |" not in content, content)
    check("code block content untouched",
          "code_line = 1" in content and T + "code_line" not in content, content)


# ------------------------------------------------------------------- excel --
def test_xlsx_bilingual_openpyxl():
    print("XLSX bilingual (openpyxl path): original + translation in cell")
    import openpyxl
    from pipeline.excel_translation_pipeline import (
        extract_excel_content_to_json, write_translated_content_to_excel)

    src = os.path.join(WORK_DIR, "bi_test.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Quarterly revenue report"
    ws["B2"] = "Net profit after taxes"
    wb.save(src)

    src_json = extract_excel_content_to_json(src, TEMP_DIR, use_xlwings=False)
    dst_json = fake_translate(src_json)
    out = write_translated_content_to_excel(src, src_json, dst_json, RESULT_DIR,
                                            use_xlwings=False, bilingual_mode=True,
                                            src_lang="en", dst_lang="ja")

    wb2 = openpyxl.load_workbook(out)
    ws2 = wb2.active
    check("A1 bilingual (original first, Excel convention)",
          ws2["A1"].value == f"Quarterly revenue report\n{T}Quarterly revenue report",
          repr(ws2["A1"].value))
    check("B2 bilingual",
          ws2["B2"].value == f"Net profit after taxes\n{T}Net profit after taxes",
          repr(ws2["B2"].value))

    # Monolingual openpyxl path unchanged
    out_mono = write_translated_content_to_excel(src, src_json, dst_json, RESULT_DIR,
                                                 use_xlwings=False, bilingual_mode=False,
                                                 src_lang="en", dst_lang="fr")
    wb3 = openpyxl.load_workbook(out_mono)
    check("monolingual path unchanged",
          wb3.active["A1"].value == T + "Quarterly revenue report",
          repr(wb3.active["A1"].value))


# -------------------------------------------------------------------- word --
def test_docx_bilingual():
    print("DOCX bilingual: original + translation, CJK brackets survive for ja target")
    from docx import Document
    from pipeline.word_translation_pipeline import (
        extract_word_content_to_json, write_translated_content_to_word)

    src = os.path.join(WORK_DIR, "bi_test.docx")
    doc = Document()
    doc.add_paragraph("これは《重要》な文書です")
    doc.add_paragraph("Second paragraph of content")
    doc.save(src)

    src_json = extract_word_content_to_json(src, TEMP_DIR)
    dst_json = fake_translate(src_json)
    out = write_translated_content_to_word(src, src_json, dst_json, TEMP_DIR, RESULT_DIR,
                                           bilingual_mode=True, src_lang="ja", dst_lang="ja")
    with zipfile.ZipFile(out) as z:
        xml = z.read("word/document.xml").decode("utf-8")

    check("original text present", "これは《重要》な文書です" in xml, xml[:500])
    check("translation present", T in xml, xml[:500])
    check("《》 NOT stripped from translation for CJK target",
          xml.count("《重要》") == 2, f"count={xml.count('《重要》')}")
    check("second paragraph bilingual",
          "Second paragraph of content" in xml and T + "Second paragraph of content" in xml, xml[:500])


# -------------------------------------------------------------------- html --
def test_html_bilingual():
    print("HTML bilingual: original inserted as sibling block after translation")
    from pipeline.html_translation_pipeline import (
        extract_html_content_to_json, write_translated_content_to_html)
    from lxml import html as lxml_html

    src = os.path.join(WORK_DIR, "bi_test.html")
    with open(src, "w", encoding="utf-8") as f:
        f.write("<!DOCTYPE html><html><head><title>Doc title text</title>"
                "<style>body { color: red; }</style></head>"
                "<body><h1>Main heading content</h1>"
                "<p>First paragraph of body text.</p>"
                "<p>Second paragraph here too.</p>"
                "</body></html>")

    src_json = extract_html_content_to_json(src, TEMP_DIR)
    dst_json = fake_translate(src_json)
    out = write_translated_content_to_html(src, src_json, dst_json, TEMP_DIR, RESULT_DIR,
                                           src_lang="en", dst_lang="ja", bilingual_mode=True)
    with open(out, encoding="utf-8") as f:
        content = f.read()

    check("translation present", T + "First paragraph of body text." in content, content)
    check("original present alongside translation",
          content.count("First paragraph of body text.") == 2, content)
    check("heading bilingual", T + "Main heading content" in content
          and content.count("Main heading content") == 2, content)
    check("css untouched", "color: red" in content, content)

    # Structure intact: each translated <p> is immediately followed by a
    # sibling <p> holding the original (translation first, original after)
    root = lxml_html.fromstring(content)
    body = root.find(".//body")
    blocks = [el for el in body if el.tag in ("h1", "p")]
    texts = [el.text_content() for el in blocks]
    idx_t = texts.index(T + "First paragraph of body text.")
    check("original is a sibling right after the translated block",
          texts[idx_t + 1] == "First paragraph of body text.", str(texts))
    check("p element count doubled (4 blocks: 2 translated + 2 original)",
          sum(1 for el in body if el.tag == "p") == 4,
          str([el.tag for el in body]))

    # Monolingual default unchanged
    out_mono = write_translated_content_to_html(src, src_json, dst_json, TEMP_DIR, RESULT_DIR,
                                                src_lang="en", dst_lang="fr")
    with open(out_mono, encoding="utf-8") as f:
        mono = f.read()
    check("default stays monolingual",
          mono.count("First paragraph of body text.") == 1, mono)


# -------------------------------------------------------------------- epub --
def test_epub_bilingual():
    print("EPUB bilingual: original sibling block after translation, zip structure intact")
    from pipeline.epub_translation_pipeline import (
        extract_epub_content_to_json, write_translated_content_to_epub)
    from lxml import etree

    src = os.path.join(WORK_DIR, "bi_book.epub")
    chapter = (
        '<?xml version="1.0" encoding="utf-8"?>'
        '<html xmlns="http://www.w3.org/1999/xhtml"><head><title>Chapter Title Text</title></head>'
        "<body><h1>The First Chapter</h1>"
        "<p>Plain narrative paragraph text.</p>"
        "<p>Another paragraph of prose.</p>"
        "</body></html>"
    )
    opf = ('<?xml version="1.0"?><package xmlns="http://www.idpf.org/2007/opf" version="2.0" '
           'unique-identifier="id"><metadata/><manifest>'
           '<item id="c1" href="chapter1.xhtml" media-type="application/xhtml+xml"/>'
           '</manifest><spine><itemref idref="c1"/></spine></package>')
    container = ('<?xml version="1.0"?><container version="1.0" '
                 'xmlns="urn:oasis:names:tc:opendocument:xmlns:container"><rootfiles>'
                 '<rootfile full-path="content.opf" media-type="application/oebps-package+xml"/>'
                 "</rootfiles></container>")
    with zipfile.ZipFile(src, "w") as z:
        z.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        z.writestr("META-INF/container.xml", container, compress_type=zipfile.ZIP_DEFLATED)
        z.writestr("content.opf", opf, compress_type=zipfile.ZIP_DEFLATED)
        z.writestr("chapter1.xhtml", chapter, compress_type=zipfile.ZIP_DEFLATED)

    src_json = extract_epub_content_to_json(src, TEMP_DIR)
    dst_json = fake_translate(src_json)
    out = write_translated_content_to_epub(src, src_json, dst_json, TEMP_DIR, RESULT_DIR,
                                           src_lang="en", dst_lang="ja", bilingual_mode=True)

    with zipfile.ZipFile(out) as z:
        first = z.infolist()[0]
        chap = z.read("chapter1.xhtml").decode("utf-8")
        names = z.namelist()

    check("mimetype first and uncompressed",
          first.filename == "mimetype" and first.compress_type == zipfile.ZIP_STORED,
          f"{first.filename} / {first.compress_type}")
    check("all members preserved", set(names) == {"mimetype", "META-INF/container.xml",
                                                  "content.opf", "chapter1.xhtml"}, str(names))
    check("translation present", T + "Plain narrative paragraph text." in chap, chap)
    check("original present alongside translation",
          chap.count("Plain narrative paragraph text.") == 2, chap)
    check("heading bilingual",
          T + "The First Chapter" in chap and chap.count("The First Chapter") == 2, chap)

    # Structure intact: original is a sibling block right after the translation
    root = etree.fromstring(chap.encode("utf-8"))
    body = next(el for el in root.iter() if etree.QName(el).localname.lower() == "body")
    blocks = [el for el in body if etree.QName(el).localname.lower() in ("h1", "p")]
    texts = ["".join(el.itertext()) for el in blocks]
    idx_t = texts.index(T + "Plain narrative paragraph text.")
    check("original is a sibling right after the translated block",
          texts[idx_t + 1] == "Plain narrative paragraph text.", str(texts))
    p_count = sum(1 for el in body if etree.QName(el).localname.lower() == "p")
    check("p element count doubled (4: 2 translated + 2 original)", p_count == 4,
          str([etree.QName(el).localname for el in body]))

    # Monolingual default unchanged
    out_mono = write_translated_content_to_epub(src, src_json, dst_json, TEMP_DIR, RESULT_DIR,
                                                src_lang="en", dst_lang="fr")
    with zipfile.ZipFile(out_mono) as z:
        chap_mono = z.read("chapter1.xhtml").decode("utf-8")
    check("default stays monolingual",
          chap_mono.count("Plain narrative paragraph text.") == 1, chap_mono)


# ------------------------------------------------------------- model fetch --
def test_model_fetch():
    print("Online model auto-fetch: /models -> (Fetched) configs, graceful failure")
    import llmWrapper.online_translation as ot

    config_dir = os.path.join("config", "api_config")
    base_name = "(Test) FetchBase"
    base_path = os.path.join(config_dir, f"{base_name}.json")
    created = [base_path]

    class FakeModelsClient:
        ids = []
        fail = False
        last_base_url = None
        last_timeout = None

        def __init__(self, api_key=None, base_url=None, timeout=None):
            FakeModelsClient.last_base_url = base_url
            FakeModelsClient.last_timeout = timeout

            class _Entry:
                def __init__(self, mid):
                    self.id = mid

            class _Models:
                @staticmethod
                def list():
                    if FakeModelsClient.fail:
                        raise RuntimeError("connection refused")

                    class _Page:
                        pass

                    page = _Page()
                    page.data = [_Entry(i) for i in FakeModelsClient.ids]
                    return page

            self.models = _Models()

    base_config = {"base_url": "https://fake.example/v1", "model": "base-model",
                   "temperature": 0.3, "top_p": 0.9}
    with open(base_path, "w", encoding="utf-8") as f:
        json.dump(base_config, f)

    original = ot.OpenAI
    ot.OpenAI = FakeModelsClient
    try:
        # 1. successful fetch writes configs for uncovered ids only
        FakeModelsClient.ids = ["base-model", "new-model-a", "weird/model:b"]
        added, error = ot.fetch_models_into_configs(base_name, "sk-test-key")
        check("two new configs written, covered id skipped", added == 2 and error is None,
              f"added={added}, error={error}")
        check("client got the selected config's base_url",
              FakeModelsClient.last_base_url == "https://fake.example/v1",
              str(FakeModelsClient.last_base_url))
        check("5s timeout applied", FakeModelsClient.last_timeout == 5,
              str(FakeModelsClient.last_timeout))

        fetched_a = os.path.join(config_dir, "(Fetched) new-model-a.json")
        fetched_b = os.path.join(config_dir, "(Fetched) weird_model_b.json")
        created.extend([fetched_a, fetched_b])

        ok_a = os.path.exists(fetched_a)
        check("(Fetched) new-model-a.json exists", ok_a)
        if ok_a:
            with open(fetched_a, encoding="utf-8") as f:
                cfg = json.load(f)
            check("fetched config copies base_url + params, swaps model",
                  cfg["model"] == "new-model-a" and cfg["base_url"] == base_config["base_url"]
                  and cfg["temperature"] == 0.3 and cfg["top_p"] == 0.9, str(cfg))
        ok_b = os.path.exists(fetched_b)
        check("filename-unsafe model id sanitized", ok_b, fetched_b)
        if ok_b:
            with open(fetched_b, encoding="utf-8") as f:
                check("sanitized file keeps real model id",
                      json.load(f)["model"] == "weird/model:b")

        # 2. re-fetch overwrites the same files (de-dupe, no extras)
        before = sorted(f for f in os.listdir(config_dir) if f.startswith("(Fetched)"))
        added2, error2 = ot.fetch_models_into_configs(base_name, "sk-test-key")
        after = sorted(f for f in os.listdir(config_dir) if f.startswith("(Fetched)"))
        check("re-fetch de-dupes (same file set)", before == after and error2 is None,
              f"{before} vs {after}")

        # 3. failed fetch is graceful: no files, error message returned
        FakeModelsClient.fail = True
        added3, error3 = ot.fetch_models_into_configs(base_name, "sk-test-key")
        check("failure returns 0 + message", added3 == 0 and error3 and "failed" in error3.lower(),
              f"added={added3}, error={error3}")
        FakeModelsClient.fail = False

        # 4. malformed selection (no such config) is graceful
        added4, error4 = ot.fetch_models_into_configs("NoSuchModelConfig", "sk-test-key")
        check("missing config returns 0 + message", added4 == 0 and error4 is not None,
              f"added={added4}, error={error4}")
    finally:
        ot.OpenAI = original
        for path in created:
            if os.path.exists(path):
                os.remove(path)


def main():
    shutil.rmtree(WORK_DIR, ignore_errors=True)
    os.makedirs(TEMP_DIR, exist_ok=True)
    os.makedirs(RESULT_DIR, exist_ok=True)

    for fn in (test_srt_bilingual, test_vtt_bilingual, test_txt_bilingual,
               test_md_bilingual, test_xlsx_bilingual_openpyxl, test_docx_bilingual,
               test_html_bilingual, test_epub_bilingual, test_model_fetch):
        try:
            fn()
        except Exception:
            import traceback
            traceback.print_exc()
            FAILED.append(fn.__name__ + " (crashed)")
        print()

    print(f"{len(PASSED)} passed, {len(FAILED)} failed")
    for name in FAILED:
        print(f"  FAIL: {name}")
    sys.exit(1 if FAILED else 0)


if __name__ == "__main__":
    main()
