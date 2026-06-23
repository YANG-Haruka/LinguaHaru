# Tests for the proofreading engine (core.backend), used by BOTH the Qt
# desktop app and the FastAPI web app:
#   1. A real translation run (translator classes with a faked LLM) writes
#      manifest.json + a copy of the original into temp/<doc>/.
#   2. save_proofread_table writes edited 'translated' values back into
#      dst_translated.json (and rejects row-count mismatches).
#   3. export_proofread_doc regenerates the document and the edited text
#      appears in it.
# Covers TXT and DOCX end-to-end, and that PDF is now listed + its export routes
# to the BabelDOC re-render path carrying the edited translations.
#
# The base temp/result/log dirs are monkeypatched to a sandbox, so there is no
# churn to config/system_config.json.
#
# Run from the repo root:
#   python tests/test_proofread.py
import json
import os
import shutil
import sys
import zipfile

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, REPO_ROOT)
os.chdir(REPO_ROOT)

# Console-safe output on Windows (CJK text)
for stream in (sys.stdout, sys.stderr):
    if hasattr(stream, "reconfigure"):
        stream.reconfigure(encoding="utf-8", errors="replace")

WORK_DIR = os.path.join(REPO_ROOT, "tests", "_proofread_work")
TEMP_DIR = os.path.join(WORK_DIR, "temp")
RESULT_DIR = os.path.join(WORK_DIR, "result")
LOG_DIR = os.path.join(WORK_DIR, "log")

T = "[T]"


def _prepare_sandbox():
    """Reset the test sandbox and point the proofread engine's storage at it.
    Used by both the __main__ runner and the pytest fixture below."""
    shutil.rmtree(WORK_DIR, ignore_errors=True)
    for d in (WORK_DIR, TEMP_DIR, RESULT_DIR, LOG_DIR):
        os.makedirs(d, exist_ok=True)
    backend.get_custom_paths = lambda: (TEMP_DIR, RESULT_DIR, LOG_DIR)


try:   # under pytest: run the same setup before each test (also makes it green there)
    import pytest

    @pytest.fixture(autouse=True)
    def _pytest_sandbox():
        _prepare_sandbox()
        yield
except ImportError:
    pass
EDIT_MARK = "EDITED-BY-PROOFREADER"

PASSED, FAILED = [], []

import core.backend as backend


def check(name, cond, detail=""):
    (PASSED if cond else FAILED).append(name)
    print(f"  [{'PASS' if cond else 'FAIL'}] {name}" + (f" -> {detail}" if detail and not cond else ""))
    return bool(cond)


def fake_translate_text(segments, previous_text, model, use_online, api_key,
                        system_prompt, user_prompt, previous_prompt, glossary_prompt,
                        glossary_terms=None, check_stop_callback=None, **kwargs):
    """Fake LLM: prefix every value with [T] and echo the same keys back."""
    from core.engine.translation_checker import clean_json
    data = json.loads(clean_json(segments))
    out = {k: T + v for k, v in data.items()}
    return (json.dumps(out, ensure_ascii=False), True,
            {"prompt_tokens": 1, "completion_tokens": 1, "total_tokens": 2})


def run_fake_translation(translator_class, src_file, ext):
    """Run a full translator-class pipeline with the faked LLM."""
    import core.engine.base_translator as bt
    original = bt.translate_text
    bt.translate_text = fake_translate_text
    try:
        translator = translator_class(
            src_file, "fake-model", False, "", "en", "fr", False,
            max_token=768, max_retries=2, thread_count=1, glossary_path=None,
            temp_dir=TEMP_DIR, result_dir=RESULT_DIR, session_lang="en", log_dir=LOG_DIR
        )
        return translator.process(os.path.splitext(src_file)[0], ext)
    finally:
        bt.translate_text = original


def proofread_cycle(doc_name, edit_row_value):
    """Load table -> edit one translated cell -> save -> re-export.

    Returns (exported_path or None)."""
    rows = [list(r) for r in backend.load_proofread_table(doc_name)]
    check(f"{doc_name}: table loaded with 3 columns",
          bool(rows) and len(rows[0]) == 3, str(rows[:1]))

    # Mismatch guard: a table with a dropped row must be rejected
    mismatch_rejected = False
    try:
        backend.save_proofread_table(doc_name, rows[:-1])
    except ValueError:
        mismatch_rejected = True
    check(f"{doc_name}: row-count mismatch is rejected", mismatch_rejected)

    # Edit the translated column (index 2) of the first row
    rows[0][2] = edit_row_value
    changed = backend.save_proofread_table(doc_name, [tuple(r) for r in rows])
    check(f"{doc_name}: edits saved", changed == 1, f"changed={changed}")

    # The edit must land in dst_translated.json (translated field only)
    dst_path = os.path.join(TEMP_DIR, doc_name, "dst_translated.json")
    with open(dst_path, encoding="utf-8") as f:
        data = json.load(f)
    check(f"{doc_name}: dst_translated.json updated",
          data[0]["translated"] == edit_row_value, repr(data[0]))
    check(f"{doc_name}: original field untouched",
          not data[0]["original"].startswith(T), repr(data[0]["original"]))

    exported = backend.export_proofread_doc(doc_name)
    ok = check(f"{doc_name}: re-export produced a file",
               exported and os.path.exists(exported), str(exported))
    return exported if ok else None


def test_txt():
    print("TXT: fake translation -> manifest -> edit -> re-export")
    from core.translators.txt_translator import TxtTranslator

    src_file = os.path.join(WORK_DIR, "proof_sample_txt.txt")
    with open(src_file, "w", encoding="utf-8") as f:
        f.write("First line of the document\n"
                "Second line with more text\n"
                "Third and final line\n")

    out_path, missing = run_fake_translation(TxtTranslator, src_file, ".txt")
    check("txt translation output exists", os.path.exists(out_path), out_path)
    check("txt no missing segments", not missing, str(missing))

    doc_dir = os.path.join(TEMP_DIR, "proof_sample_txt")
    manifest_path = os.path.join(doc_dir, "manifest.json")
    check("manifest.json written", os.path.exists(manifest_path), doc_dir)
    with open(manifest_path, encoding="utf-8") as f:
        manifest = json.load(f)
    check("manifest fields correct",
          manifest.get("file_extension") == ".txt" and manifest.get("src_lang") == "en"
          and manifest.get("dst_lang") == "fr" and manifest.get("model") == "fake-model"
          and manifest.get("input_file") == "proof_sample_txt.txt", str(manifest))
    original_copy = os.path.join(doc_dir, manifest.get("original_copy", ""))
    check("original file copied into temp folder",
          os.path.exists(original_copy)
          and open(original_copy, encoding="utf-8").read() == open(src_file, encoding="utf-8").read(),
          original_copy)

    docs = backend.list_proofread_docs()
    check("doc listed for proofreading", "proof_sample_txt" in docs, str(docs))

    exported = proofread_cycle("proof_sample_txt", EDIT_MARK + " ligne une")
    if exported:
        with open(exported, encoding="utf-8") as f:
            content = f.read()
        check("edited text appears in regenerated txt", EDIT_MARK in content, content)
        check("untouched rows keep the fake translation",
              T + "Second line with more text" in content, content)


def test_docx():
    print("DOCX: fake translation -> manifest -> edit -> re-export")
    from docx import Document
    from core.translators.word_translator import WordTranslator

    src_file = os.path.join(WORK_DIR, "proof_sample_docx.docx")
    doc = Document()
    doc.add_paragraph("Opening paragraph for the proofread test")
    doc.add_paragraph("Closing paragraph with different words")
    doc.save(src_file)

    out_path, missing = run_fake_translation(WordTranslator, src_file, ".docx")
    check("docx translation output exists", os.path.exists(out_path), out_path)

    doc_dir = os.path.join(TEMP_DIR, "proof_sample_docx")
    check("docx manifest written", os.path.exists(os.path.join(doc_dir, "manifest.json")), doc_dir)
    check("docx listed for proofreading", "proof_sample_docx" in backend.list_proofread_docs(),
          str(backend.list_proofread_docs()))

    exported = proofread_cycle("proof_sample_docx", EDIT_MARK + " premiere phrase")
    if exported:
        with zipfile.ZipFile(exported) as z:
            xml = z.read("word/document.xml").decode("utf-8")
        check("edited text appears in regenerated docx", EDIT_MARK in xml, xml[:400])
        check("untouched paragraph keeps the fake translation",
              T + "Closing paragraph with different words" in xml, xml[:400])


def test_pdf_included():
    print("PDF: listed for proofreading + export routes to a BabelDOC re-render")
    pdf_dir = os.path.join(TEMP_DIR, "fake_pdf_doc")
    os.makedirs(pdf_dir, exist_ok=True)
    rows = [{"count_src": 1, "original": "Hello world", "translated": T + "Hello world"},
            {"count_src": 2, "original": "Second para", "translated": T + "Second para"}]
    with open(os.path.join(pdf_dir, "dst_translated.json"), "w", encoding="utf-8") as f:
        json.dump(rows, f)
    with open(os.path.join(pdf_dir, "fake_pdf_doc.pdf"), "wb") as f:
        f.write(b"%PDF-1.4 fake")
    with open(os.path.join(pdf_dir, "manifest.json"), "w", encoding="utf-8") as f:
        json.dump({"file_extension": ".pdf", "original_copy": "fake_pdf_doc.pdf",
                   "src_lang": "en", "dst_lang": "fr", "model": "fake",
                   "bilingual_mode": False}, f)

    check("pdf doc now listed", "fake_pdf_doc" in backend.list_proofread_docs(),
          str(backend.list_proofread_docs()))

    table = backend.load_proofread_table("fake_pdf_doc")
    check("pdf table loaded (2 rows)", len(table) == 2, str(table))
    edited = [list(r) for r in table]
    edited[0][2] = EDIT_MARK + " bonjour"
    changed = backend.save_proofread_table("fake_pdf_doc", [tuple(r) for r in edited])
    check("pdf edit saved", changed == 1, f"changed={changed}")

    # Export must route PDF to the BabelDOC re-render path, passing the EDITED
    # translations as overrides. Stub the heavy re-render (no BabelDOC needed).
    captured = {}

    def fake_pdf_export(folder, manifest, doc_name, dst_json, original_copy):
        with open(dst_json, encoding="utf-8") as f:
            data = json.load(f)
        captured["overrides"] = {it["original"]: it["translated"] for it in data}
        out = os.path.join(RESULT_DIR, "fake_pdf_doc_en2fr_proofread.pdf")
        with open(out, "wb") as f:
            f.write(b"%PDF re-rendered")
        return out

    orig = backend._export_pdf_proofread
    backend._export_pdf_proofread = fake_pdf_export
    try:
        exported = backend.export_proofread_doc("fake_pdf_doc")
    finally:
        backend._export_pdf_proofread = orig
    check("pdf export routed to BabelDOC re-render",
          bool(exported) and os.path.exists(exported), str(exported))
    check("edited translation handed to the re-render",
          captured.get("overrides", {}).get("Hello world") == EDIT_MARK + " bonjour",
          str(captured.get("overrides")))
    shutil.rmtree(pdf_dir, ignore_errors=True)


def main():
    _prepare_sandbox()

    # NOTE: a fresh translation clears the whole temp dir, so each format must
    # finish its full proofread cycle before the next one starts.
    for fn in (test_txt, test_docx, test_pdf_included):
        try:
            fn()
        except Exception:
            import traceback
            traceback.print_exc()
            FAILED.append(fn.__name__ + " (crashed)")
        print()

    print("=" * 60)
    print(f"{len(PASSED)} passed, {len(FAILED)} failed")
    for name in FAILED:
        print(f"  FAIL: {name}")
    sys.exit(1 if FAILED else 0)


if __name__ == "__main__":
    main()
