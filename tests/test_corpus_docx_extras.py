# Corpus tests: DOCX comments, endnotes and charts translation.
#
# python-docx can't author comments/endnotes/charts, so a base doc is built
# and those parts are injected into the .docx zip, then run through the real
# extract -> translate -> write round trip.
#
# Run from the repo root:
#   python tests/test_corpus_docx_extras.py
import os
import shutil
import sys
import zipfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from tests.corpus_common import T, check, fake_translate, run, work_dirs

WORK_DIR, TEMP_DIR, RESULT_DIR = work_dirs("docx_extras")

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

COMMENTS = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    f'<w:comments xmlns:w="{W}">'
    '<w:comment w:id="0" w:author="Rev" w:date="2024-01-01T00:00:00Z" w:initials="R">'
    "<w:p><w:r><w:t>Comment text to translate</w:t></w:r></w:p></w:comment></w:comments>")

ENDNOTES = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    f'<w:endnotes xmlns:w="{W}">'
    '<w:endnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:endnote>'
    '<w:endnote w:type="continuationSeparator" w:id="0">'
    "<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:endnote>"
    '<w:endnote w:id="1"><w:p><w:r><w:t>Endnote text to translate</w:t></w:r></w:p></w:endnote>'
    "</w:endnotes>")

CHART = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart"'
    ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><c:chart>'
    '<c:title><c:tx><c:rich><a:p><a:r><a:t>Sales Chart Title</a:t></a:r></a:p>'
    '</c:rich></c:tx></c:title><c:plotArea><c:barChart><c:ser><c:cat><c:strRef>'
    '<c:f>Sheet1!$A$1</c:f><c:strCache><c:ptCount val="1"/><c:pt idx="0">'
    '<c:v>Quarter One Label</c:v></c:pt></c:strCache></c:strRef></c:cat>'
    "</c:ser></c:barChart></c:plotArea></c:chart></c:chartSpace>")

OVERRIDES = (
    b'<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-'
    b'officedocument.wordprocessingml.comments+xml"/>'
    b'<Override PartName="/word/endnotes.xml" ContentType="application/vnd.openxmlformats-'
    b'officedocument.wordprocessingml.endnotes+xml"/>'
    b'<Override PartName="/word/charts/chart1.xml" ContentType="application/vnd.openxmlformats-'
    b'officedocument.drawingml.chart+xml"/>')


def _build_docx_with_parts(path):
    from docx import Document
    base = path + ".base"
    doc = Document()
    doc.add_paragraph("Main body paragraph text.")
    doc.save(base)

    with zipfile.ZipFile(base) as zin, zipfile.ZipFile(path, "w") as zout:
        for n in zin.namelist():
            data = zin.read(n)
            if n == "[Content_Types].xml":
                data = data.replace(b"</Types>", OVERRIDES + b"</Types>")
            zout.writestr(n, data)
        zout.writestr("word/comments.xml", COMMENTS)
        zout.writestr("word/endnotes.xml", ENDNOTES)
        zout.writestr("word/charts/chart1.xml", CHART)
    os.remove(base)


def _png_bytes():
    import io
    from PIL import Image
    buf = io.BytesIO()
    Image.new("RGB", (2, 2), (200, 100, 50)).save(buf, format="PNG")
    return buf.getvalue()


def _build_docx_with_alttext(path):
    """A docx with one inline picture carrying descr/title alt text on wp:docPr.

    python-docx authors the drawing; the accessibility attributes live as
    attributes on the drawing property elements, which we set on the lxml tree.
    """
    import io
    from docx import Document

    doc = Document()
    doc.add_paragraph("Body text before image.")
    run = doc.add_paragraph().add_run()
    run.add_picture(io.BytesIO(_png_bytes()))

    WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    body = doc.element.body
    docprs = body.findall(f".//{{{WP}}}docPr")
    check("docx authored a wp:docPr", len(docprs) >= 1, str(len(docprs)))
    docprs[0].set("descr", "Description alt text")
    docprs[0].set("title", "Title alt text")
    cnvprs = body.findall(f".//{{{PIC}}}cNvPr")
    if cnvprs:
        cnvprs[0].set("descr", "Pic description alt text")
    doc.save(path)


def test_docx_drawing_alttext():
    print("DOCX: drawing alt-text (docPr/cNvPr descr & title) translated")
    from core.pipelines.word_translation_pipeline import (
        extract_word_content_to_json, write_translated_content_to_word)

    src = os.path.join(WORK_DIR, "alttext.docx")
    _build_docx_with_alttext(src)

    src_json = extract_word_content_to_json(src, TEMP_DIR)
    import json
    with open(src_json, encoding="utf-8") as f:
        items = json.load(f)
    values = [i["value"] for i in items]
    alttext_items = [i for i in items if i.get("type") == "word_alttext"]

    check("docPr descr extracted", "Description alt text" in values, str(values))
    check("docPr title extracted", "Title alt text" in values, str(values))
    check("cNvPr descr extracted", "Pic description alt text" in values, str(values))
    check("alttext items carry part/attr",
          all("part" in i and "attr" in i and "elem_kind" in i for i in alttext_items),
          str(alttext_items))

    dst_json = fake_translate(src_json)
    out = write_translated_content_to_word(src, src_json, dst_json, TEMP_DIR, RESULT_DIR,
                                           bilingual_mode=False, src_lang="en", dst_lang="ja")

    with zipfile.ZipFile(out) as z:
        body = z.read("word/document.xml").decode("utf-8")

    check("docPr descr translated", T + "Description alt text" in body, body)
    check("docPr title translated", T + "Title alt text" in body, body)
    check("cNvPr descr translated", T + "Pic description alt text" in body, body)
    check("drawing element survived", "<w:drawing" in body, body)
    check("image blip survived", "blip" in body, body)
    check("body paragraph still translated", T + "Body text before image." in body, body)


def test_docx_comments_endnotes_charts():
    print("DOCX: comments, endnotes and chart text translated")
    from core.pipelines.word_translation_pipeline import (
        extract_word_content_to_json, write_translated_content_to_word)

    src = os.path.join(WORK_DIR, "extras.docx")
    _build_docx_with_parts(src)

    src_json = extract_word_content_to_json(src, TEMP_DIR)
    import json
    with open(src_json, encoding="utf-8") as f:
        extracted = [i["value"] for i in json.load(f)]
    check("comment text extracted", "Comment text to translate" in extracted, str(extracted))
    check("endnote text extracted", "Endnote text to translate" in extracted, str(extracted))
    check("chart title extracted", "Sales Chart Title" in extracted, str(extracted))
    check("chart category cache extracted", "Quarter One Label" in extracted, str(extracted))

    dst_json = fake_translate(src_json)
    out = write_translated_content_to_word(src, src_json, dst_json, TEMP_DIR, RESULT_DIR,
                                           bilingual_mode=False, src_lang="en", dst_lang="ja")

    with zipfile.ZipFile(out) as z:
        comments = z.read("word/comments.xml").decode("utf-8")
        endnotes = z.read("word/endnotes.xml").decode("utf-8")
        chart = z.read("word/charts/chart1.xml").decode("utf-8")
        body = z.read("word/document.xml").decode("utf-8")

    check("comment translated", T + "Comment text to translate" in comments, comments)
    check("endnote translated", T + "Endnote text to translate" in endnotes, endnotes)
    check("endnote separators untouched", "<w:separator/>" in endnotes, endnotes)
    check("chart title translated", T + "Sales Chart Title" in chart, chart)
    check("chart category cache translated", T + "Quarter One Label" in chart, chart)
    check("body paragraph still translated", T + "Main body paragraph text." in body, body)


def test_docx_mixed_format_runs_preserved():
    print("DOCX: mixed-format runs preserved (proportional distribution, no collapse)")
    from docx import Document
    from lxml import etree
    from core.pipelines.word_translation_pipeline import (
        extract_word_content_to_json, write_translated_content_to_word)

    src = os.path.join(WORK_DIR, "mixed_runs.docx")
    doc = Document()
    # One paragraph, three runs with DISTINCT formatting: normal / bold / italic.
    p = doc.add_paragraph()
    p.add_run("Normal text ")
    p.add_run("bold text ").bold = True
    p.add_run("italic text").italic = True
    original_joined = "Normal text bold text italic text"
    doc.save(src)

    src_json = extract_word_content_to_json(src, TEMP_DIR)
    dst_json = fake_translate(src_json)
    out = write_translated_content_to_word(src, src_json, dst_json, TEMP_DIR, RESULT_DIR,
                                           bilingual_mode=False, src_lang="en", dst_lang="ja")

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(out) as z:
        tree = etree.fromstring(z.read("word/document.xml"))

    ns = {"w": W_NS}
    # Locate the body paragraph carrying our text.
    target_p = None
    for para in tree.iter(f"{{{W_NS}}}p"):
        full = "".join(t.text or "" for t in para.iter(f"{{{W_NS}}}t"))
        if "Normal text" in full and "italic text" in full:
            target_p = para
            break

    check("mixed-run paragraph found in output", target_p is not None)
    if target_p is None:
        return

    full_text = "".join(t.text or "" for t in target_p.iter(f"{{{W_NS}}}t"))
    # (a) no text lost or duplicated: full paragraph text == [T] + original join
    check("full paragraph text equals [T] + original (no loss/dup)",
          full_text == T + original_joined, full_text)

    runs = target_p.xpath("./w:r", namespaces=ns)
    # (b) formatting not collapsed to a single run
    check("multiple runs preserved (not collapsed)", len(runs) > 1, f"run count={len(runs)}")

    has_bold = any(r.xpath("./w:rPr/w:b", namespaces=ns) for r in runs)
    has_italic = any(r.xpath("./w:rPr/w:i", namespaces=ns) for r in runs)
    # (c) at least one run keeps bold and one keeps italic (rPr preserved)
    check("a run still carries w:b (bold rPr preserved)", has_bold,
          etree.tostring(target_p).decode("utf-8"))
    check("a run still carries w:i (italic rPr preserved)", has_italic,
          etree.tostring(target_p).decode("utf-8"))


if __name__ == "__main__":
    run([test_docx_drawing_alttext, test_docx_comments_endnotes_charts,
         test_docx_mixed_format_runs_preserved])
